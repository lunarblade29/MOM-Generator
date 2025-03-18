from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
import os
import re
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement

app = Flask(__name__)

# Path to your Word document template
TEMPLATE_PATH = os.path.join('templates', 'MOM_template.docx')


# ✅ Extract meeting date and agenda points from email content

def extract_meeting_details(text):
    # Extract raw date from email text
    date_match = re.search(r"scheduled on ([A-Za-z]+, [A-Za-z]+ \d{1,2}, \d{4})", text)
    if date_match:
        raw_date = date_match.group(1)
        try:
            # Remove weekday (e.g., "Wednesday,") and parse the rest
            clean_date = " ".join(raw_date.split()[1:])  # e.g., 'November 01, 2023'
            parsed_date = datetime.strptime(clean_date, "%B %d, %Y")
            meeting_date = parsed_date.strftime("%B %d, %Y")  # Format: December 15, 2023
        except Exception:
            meeting_date = raw_date  # fallback
    else:
        meeting_date = "Unknown Date"

    # Extract agenda section
    agenda_section = re.search(r"(1\..*?)(?:\n\n|\Z)", text, re.DOTALL)
    agenda_text = agenda_section.group(1).strip() if agenda_section else ""

    # Extract individual agenda items
    agenda_points = re.findall(r"\d+\..*", agenda_text)

    return meeting_date, agenda_points


# ✅ Replace placeholder text in header
def replace_header_text(doc, placeholder, replacement):
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)


# ✅ Proper XML-based paragraph insertion after a reference paragraph
def insert_paragraph_after_paragraph(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = paragraph._parent.add_paragraph(text)
    if style:
        new_para.style = style
    return new_para


# ✅ Generate final Word document with inserted agenda

def generate_doc(meeting_date, agenda_points):
    doc = Document("templates/MOM_template.docx")

    # Replace meeting date
    for paragraph in doc.paragraphs:
        if "<<MEETING_DATE>>" in paragraph.text:
            replaced_text = paragraph.text.replace("<<MEETING_DATE>>", meeting_date)
            paragraph.clear()  # Clear the paragraph text and runs
            run = paragraph.add_run(replaced_text)

    # Replace meeting date
    for paragraph in doc.paragraphs:
        if "<<MEETING_DATE_u>>" in paragraph.text:
            replaced_text = paragraph.text.replace("<<MEETING_DATE_u>>", meeting_date)
            paragraph.clear()  # Clear the paragraph text and runs
            run = paragraph.add_run(replaced_text)
            run.underline = True

    # Replace in header
    replace_header_text(doc, "<<MEETING_DATE>>", meeting_date)

    # Replace AGENDA and make agenda points bold
    for paragraph in doc.paragraphs:
        if "<<AGENDA>>" in paragraph.text:
            for agenda in agenda_points:  # insert in reverse to preserve order
                new_para = paragraph.insert_paragraph_before("")
                run = new_para.add_run(agenda)
                run.bold = True
            # Remove the placeholder
            p = paragraph._element
            p.getparent().remove(p)
            break

    output_path = f"MoM of the MBA Committee meeting dated {meeting_date}.docx"
    doc.save(output_path)
    return output_path


# ✅ Flask Routes
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        email_text = request.form['email_text']
        meeting_date, agenda_points = extract_meeting_details(email_text)
        file_path = generate_doc(meeting_date, agenda_points)
        return send_file(file_path, as_attachment=True)

    return render_template("index.html")


# ✅ Run the Flask App
if __name__ == "__main__":
    app.run(debug=True)
